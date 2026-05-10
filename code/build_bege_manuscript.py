from __future__ import annotations

import csv
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PKG = ROOT / "computational_package"
DATA = PKG / "data"
FIG = PKG / "figures"
OUT_DOCX = ROOT / "Time dependent stability of.docx"
COVER = ROOT / "Cover letter.docx"
DECL = ROOT / "Submission declarations.docx"


AUTHOR = "Gabriel Jesús Montúfar Chiriboga"
AFFIL = "Universidad de Panamá, Facultad de Ingeniería, Departamento de Ingeniería Civil, Panamá"
EMAIL = "gabriel.montufar@up.ac.pa"
ORCID = "https://orcid.org/0000-0003-3392-3728"
JOURNAL = "Bulletin of Engineering Geology and the Environment"
GITHUB_URL = "https://github.com/gabrielmontufar/article-120-slope-state-envelope"


def add_code_to_docx(source: Path, dest: Path) -> None:
    shutil.copy2(source, dest)
    with zipfile.ZipFile(dest, "a", zipfile.ZIP_DEFLATED) as z:
        z.write(source, "embedded_source/original_spanish_manuscript.docx")


def set_cell_shading(cell, fill="FFFFFF"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color="000000", size="4"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def clear_paragraph_borders(paragraph_or_style):
    ppr = paragraph_or_style._element.get_or_add_pPr()
    border = ppr.find(qn("w:pBdr"))
    if border is not None:
        ppr.remove(border)


def plain_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.size = Pt(9)
        set_cell_shading(hdr[i], "FFFFFF")
        set_cell_borders(hdr[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = str(text)
            set_cell_shading(cells[i], "FFFFFF")
            set_cell_borders(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.size = Pt(9)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    return table


def p(doc: Document, text: str = "", style: str | None = None, align=None):
    para = doc.add_paragraph(style=style)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing = 1.08
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return para


def reference_paragraph(doc: Document, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.line_spacing = 1.0
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return para


def heading(doc: Document, text: str, level: int = 1):
    para = doc.add_paragraph(style=f"Heading {level}")
    para.paragraph_format.space_before = Pt(10 if level == 1 else 8)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.bold = True
    return para


def caption(doc: Document, text: str):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(10)
    return para


def read_csv(name: str):
    with (DATA / name).open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def fmt(x, nd=3):
    try:
        return f"{float(x):.{nd}f}"
    except Exception:
        return str(x)


def configure(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    styles = doc.styles
    for name in ["Normal", "Body Text"]:
        if name in styles:
            st = styles[name]
            st.font.name = "Times New Roman"
            st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            st.font.size = Pt(10)
            st.font.color.rgb = RGBColor(0, 0, 0)
    for level, size in [(1, 13), (2, 11), (3, 10)]:
        st = styles[f"Heading {level}"]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
    st = styles["Title"]
    st.font.name = "Times New Roman"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    st.font.size = Pt(16)
    st.font.bold = True
    st.font.color.rgb = RGBColor(0, 0, 0)
    clear_paragraph_borders(st)


def add_field_code(paragraph, code: str):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = code
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_text)
    run._r.append(fld_end)


def add_equation_placeholder(doc: Document, label: str, formula: str):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"{label}  {formula}")
    run.font.name = "Cambria Math"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return para


def build_manuscript():
    doc = Document()
    configure(doc)
    title = "Time-dependent stability of rainfall-infiltrated cracked, vegetated and infrastructure-modified slopes"
    title_p = doc.add_paragraph(style="Title")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    clear_paragraph_borders(title_p)
    title_p.add_run(title).font.color.rgb = RGBColor(0, 0, 0)
    p(doc, AUTHOR, align=WD_ALIGN_PARAGRAPH.CENTER)
    p(doc, f"{AFFIL}; corresponding author e-mail: {EMAIL}; ORCID: {ORCID}", align=WD_ALIGN_PARAGRAPH.CENTER)

    heading(doc, "Abstract", 1)
    p(
        doc,
        "Rainfall-induced slope instability is commonly reduced to a single factor of safety or to an empirical rainfall threshold, even though slope response depends on transient suction loss, positive pore pressure, crack filling, vegetation and infrastructure loading. This paper proposes a reproducible state-conditioned framework for time-dependent stability assessment of cracked, vegetated and infrastructure-modified slopes. The framework combines a reduced transient infiltration state, unsaturated shear strength, crack water loading, apparent root cohesion, road surcharge and drainage condition into deterministic FSmin(t), Monte Carlo Pf(t), a hydro-mechanical memory index, a crack-root compensation factor and a state-conditioned rainfall threshold envelope. A synthetic but fully reproducible demonstration with 30,000 Monte Carlo samples shows that a cracked slope reaches a cumulative failure probability of 0.669, whereas a comparable cracked vegetated slope reaches 0.043. When a road surcharge and clogged drainage are added, the probability rises to 0.883; maintained drainage and vegetation reduce it to 0.017. These results do not represent a site calibration, but they show how civil infrastructure state can be inserted explicitly into rainfall-stability decisions. The contribution is a transparent decision framework that links hydrological forcing, biological reinforcement, crack hydraulics and infrastructure maintenance to time-dependent reliability rather than to a static safety factor.",
    )
    p(doc, "Keywords: rainfall-induced landslides; unsaturated slope stability; cracks; vegetation; road surcharge; drainage; reliability")

    heading(doc, "1 Introduction", 1)
    p(
        doc,
        "Rainfall-induced shallow landslides remain a recurring engineering-geology problem because the controlling mechanism is not rainfall alone, but the time-dependent conversion of rainfall into suction loss, positive pore pressure, crack water loading and mechanical demand. Classical rainfall thresholds (Caine 1980; Guzzetti et al. 2008), physically based landslide initiation models (Montgomery and Dietrich 1994; Iverson 2000) and grid-based transient tools such as TRIGRS (Baum et al. 2008) have made these processes operational, yet many practical assessments still end in a single factor of safety. That scalar output is convenient, but it hides whether the critical state is short-lived, persistent after the storm, controlled by crack filling, or driven by an infrastructure condition such as a road surcharge or clogged drainage.",
    )
    p(
        doc,
        "The engineering-geology context is broader than natural slopes. Roads, cuts, berms, culverts, shoulder drains and embankment surcharges often alter the hydrological boundary condition and the mechanical stress state at the same time. The Bulletin of Engineering Geology and the Environment scope explicitly includes hydrological and mechanical behaviour of soil and rock masses, time-dependent property changes and stability parameters for earth masses; therefore, the relevant question is not only whether a vegetated or cracked slope is stable, but whether its time-dependent reliability changes when civil infrastructure is added to the system.",
    )
    p(
        doc,
        "Unsaturated soil mechanics provides the physical basis for this question. Richards' equation (Richards 1931), the van Genuchten-Mualem retention and conductivity relations (Mualem 1976; van Genuchten 1980), and the unsaturated shear-strength framework of Fredlund et al. (1978) explain why rainfall can degrade stability before full saturation is reached. Numerical studies of transient seepage in slopes (Ng and Shi 1998; Collins and Znidarcic 2004; Griffiths and Lu 2005; Rahardjo et al. 2008) further show that antecedent conditions and hydraulic parameters control the timing of instability.",
    )
    p(
        doc,
        "Vegetation and roots complicate the same problem. Roots can reinforce soil mechanically and modify the near-surface water balance (Waldron and Dakessian 1981; Wu et al. 1979; Gray and Sotir 1992; Morgan and Rickson 2003), but the stabilizing effect is conditional on root depth, density, species and the location of the potential failure surface (Schwarz et al. 2010; Sidle and Bogaard 2016). Similarly, reliability-based geotechnical design emphasizes that uncertain hydraulic and mechanical parameters should be propagated explicitly rather than hidden inside a deterministic factor (Duncan 2000; Phoon and Kulhawy 1999; Ang and Tang 2007; Ditlevsen and Madsen 1996; Fenton and Griffiths 2008; Melchers and Beck 2018).",
    )
    p(
        doc,
        "This paper contributes a state-conditioned temporal reliability framework that treats cracks, vegetation and infrastructure condition as explicit slope states. The novelty is not another isolated factor-of-safety equation. The proposed contribution is a compact decision architecture that produces three additional quantities: a hydro-mechanical memory index (HMMI), a crack-root compensation factor (CRCF), and a rainfall threshold envelope conditioned by biological and infrastructure state. These quantities are intended to help compare maintenance, vegetation and drainage alternatives under identical rainfall forcing.",
    )

    heading(doc, "2 Proposed state-conditioned framework", 1)
    p(
        doc,
        "The method starts from a slope state vector S={C,V,I}, where C describes crack condition, V describes vegetation/root condition and I describes civil-infrastructure condition. Infrastructure is represented by two mechanisms: surcharge at or near the crest, and drainage/runoff modification. The state vector is combined with uncertain geotechnical and hydrological parameters X to compute time-dependent infiltration, shear strength, FSmin(t) and Pf(t). Figure 1 summarizes the workflow before the mathematical details are introduced.",
    )
    caption(doc, "Fig. 1 Conceptual workflow linking rainfall, infiltration, crack-root state, road/drainage condition and temporal reliability.")
    doc.add_picture(str(FIG / "Fig1_conceptual_workflow.png"), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p(
        doc,
        "The water state is represented by a reduced infiltration memory variable W(t). This reduction is not proposed as a replacement for a full Richards solution; rather, it is a reproducible screening layer that preserves the essential temporal effect needed for reliability calculations. The form follows the same physical direction as transient infiltration models (Richards 1931; Iverson 2000; Baum et al. 2008): rainfall increases wetness, drainage dissipates it, cracks accelerate preferential inflow and vegetation can intercept rainfall and promote drying.",
    )
    add_equation_placeholder(doc, "(1)", "dW/dt = qR(t,S)/[Ks+2] - W/Td(S)")
    p(
        doc,
        "where R(t,S) is the effective rainfall after vegetation interception and infrastructure runoff concentration, Ks is saturated hydraulic conductivity and Td(S) is a drainage time scale modified by maintenance condition. A clogged drain is represented as larger runoff concentration and slower dissipation; maintained drainage is represented as reduced effective inflow and faster dissipation.",
    )
    add_equation_placeholder(doc, "(2)", "s(t)=s0 exp(-aW) + sV(t);    u+(t)=max[0,b(W-W0)] + uc(t)")
    p(
        doc,
        "The suction s(t) and positive pore pressure u+(t) are therefore functions of wetting history, not only instantaneous rainfall. This formulation supports the proposed HMMI, which measures how long a slope remains near or below a target safety state after the storm.",
    )
    add_equation_placeholder(doc, "(3)", "HMMI = integral max[0, FSref - FSmin(t)] dt")
    p(
        doc,
        "The infinite-slope stability expression is used as the transparent demonstration layer. It is not the only possible limit-equilibrium model; a slice method or finite-element strength-reduction step could replace it in later work. The present form is sufficient to expose the effect of suction, cracks, roots and infrastructure without hiding assumptions inside a black-box solver (Duncan et al. 2014; Abramson et al. 2002).",
    )
    add_equation_placeholder(doc, "(4)", "FS(t) = [c' + cr + (sigma_n + q cos^2 beta - u+) tan phi' + chi s tan phi_b] / [tau_g + q sin beta cos beta + tau_c]")
    p(
        doc,
        "Here q is the infrastructure surcharge, cr is apparent root cohesion, beta is slope angle, tau_c is crack-water driving shear, phi_b controls suction contribution and chi is an effective stress parameter. The state-dependent probability of failure is estimated by Monte Carlo as shown in Eq. (5), consistent with geotechnical reliability practice (Baecher and Christian 2003; Fenton and Griffiths 2008).",
    )
    add_equation_placeholder(doc, "(5)", "Pf(t|S) = P[FSmin(t,X,S) < 1]")
    p(
        doc,
        "The CRCF is introduced to quantify how much of the crack-induced risk is neutralized by the selected vegetation state under identical rainfall and soil uncertainty. Values near zero imply that vegetation does not compensate for crack effects; values near one imply that most of the crack-induced risk increment is offset.",
    )
    add_equation_placeholder(doc, "(6)", "CRCF = [Pf,cracked - Pf,cracked+vegetated] / [Pf,cracked - Pf,bare]")
    p(
        doc,
        "For infrastructure, a parallel drainage-resilience metric is used. It compares a road surcharge with clogged drainage against a road surcharge with maintained drainage and vegetation. The indicator is not a universal design factor; it is a transparent decision metric for ranking maintenance alternatives under the same probabilistic assumptions.",
    )
    add_equation_placeholder(doc, "(7)", "IDR = [Pf,road clogged - Pf,road maintained] / Pf,road clogged")

    heading(doc, "3 Reproducible numerical demonstration", 1)
    p(
        doc,
        "A synthetic slope was used to demonstrate the framework. The case is deliberately transparent and is not presented as a calibrated field site. The geometry and parameter ranges were selected to be plausible for a shallow residual-soil slope affected by tropical rainfall, unsaturated suction loss and shallow infrastructure loading. Table 1 lists the random variables, while Table 2 defines the six slope states evaluated.",
    )
    doc.add_page_break()
    caption(doc, "Table 1 Random variables used in the reproducible demonstration.")
    plain_table(
        doc,
        ["Variable", "Distribution or range", "Role"],
        [
            ["Slope angle beta", "Normal 32 deg, sd 2 deg, truncated 26-39 deg", "Mechanical demand"],
            ["Failure depth z", "Normal 2.10 m, sd 0.22 m, truncated 1.50-2.80 m", "Potential failure plane"],
            ["Effective cohesion c'", "Lognormal mean 8 kPa, COV about 0.25", "Shear strength"],
            ["Friction angle phi'", "Normal 33 deg, sd 2 deg, truncated 27-39 deg", "Shear strength"],
            ["Suction-friction angle phi_b", "Normal 16 deg, sd 1.8 deg, truncated 10-22 deg", "Unsaturated strength"],
            ["Initial suction s0", "Lognormal mean 32 kPa, COV about 0.32", "Hydrological memory"],
            ["Hydraulic conductivity Ks", "Lognormal mean 7.5 mm/h, COV about 0.42", "Infiltration rate"],
            ["Root cohesion cr", "Lognormal mean 5 kPa, COV about 0.40", "Vegetation reinforcement"],
            ["Crack depth", "Normal 0.75 m, sd 0.15 m, truncated 0.25-1.15 m", "Preferential flow and crack water load"],
        ],
        [1.6, 2.7, 2.1],
    )

    caption(doc, "Table 2 Slope states evaluated in the numerical demonstration.")
    plain_table(
        doc,
        ["State", "Crack", "Vegetation", "Infrastructure condition"],
        [
            ["Bare slope", "No", "No", "No surcharge; baseline drainage"],
            ["Cracked slope", "Yes", "No", "No surcharge; crack preferential inflow"],
            ["Vegetated slope", "No", "Yes", "Interception and root cohesion"],
            ["Cracked vegetated slope", "Yes", "Yes", "Crack plus vegetation effects"],
            ["Road surcharge and clogged drain", "Yes", "No", "10 kPa road surcharge; runoff concentration; poor drainage"],
            ["Road surcharge with maintained drainage", "Yes", "Yes", "10 kPa road surcharge; maintained drainage; vegetation"],
        ],
        [2.0, 1.0, 1.2, 2.4],
    )
    p(
        doc,
        "The storm consisted of a 96 h record with a peak intensity of 24 mm/h and post-peak recession. The time step was 0.25 h. All calculations were run with a fixed random seed of 120 and 30,000 Monte Carlo samples. The code, CSV files and high-resolution figures are included in the supplementary computational package.",
    )

    heading(doc, "4 Results", 1)
    p(
        doc,
        "Figure 2 shows the hydrological response produced by the common rainfall event. Cracks accelerate wetting and increase positive pore-pressure response, whereas vegetation reduces the wetness peak and partially restores suction. Infrastructure changes the response by concentrating runoff when drainage is clogged or by reducing the effective hydrological load when drainage is maintained.",
    )
    caption(doc, "Fig. 2 Rainfall forcing, suction response and positive pore-pressure response for the evaluated slope states.")
    doc.add_picture(str(FIG / "Fig2_hydrological_response.png"), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    det = read_csv("deterministic_scenario_summary.csv")
    caption(doc, "Table 3 Deterministic response summary.")
    plain_table(
        doc,
        ["Scenario", "Minimum FS", "Critical time (h)", "HMMI (FS-h)", "Maximum pore pressure (kPa)"],
        [[r["scenario"], fmt(r["min_fs"]), fmt(r["time_of_min_fs_h"], 2), fmt(r["hmmi_fs_hour"]), fmt(r["max_pore_pressure_kpa"])] for r in det],
        [2.25, 1.0, 1.0, 1.1, 1.35],
    )
    p(
        doc,
        "Table 3 indicates that the deterministic bare slope remains above unity, while the cracked slope drops to FS=0.953 and the road-surcharge state with clogged drainage drops to FS=0.789. The HMMI captures the persistence of this degradation: the road/clogged-drain state accumulates 11.882 FS-h of below-reference safety, compared with 3.957 FS-h for the cracked slope and zero for the vegetated states.",
    )
    p(
        doc,
        "Figure 3 converts this deterministic response into temporal reliability. The cracked state has a much larger Pf(t) than the bare state, but vegetation reduces both the peak probability and the duration of elevated risk. The infrastructure result is more disruptive from a practical standpoint: a road surcharge is not necessarily destabilizing by itself, but the coupled road-surcharge and clogged-drainage state becomes the dominant risk condition.",
    )
    caption(doc, "Fig. 3 Deterministic FS(t) and Monte Carlo Pf(t) for the six slope states.")
    doc.add_picture(str(FIG / "Fig3_temporal_reliability.png"), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    mc = read_csv("monte_carlo_summary.csv")
    caption(doc, "Table 4 Monte Carlo reliability summary based on 30,000 samples.")
    plain_table(
        doc,
        ["Scenario", "Peak Pf", "Cumulative Pf", "Median min FS", "5th pct. min FS", "Reliability index"],
        [[r["scenario"], fmt(r["peak_pf"], 4), fmt(r["cumulative_pf"], 4), fmt(r["median_min_fs"]), fmt(r["p05_min_fs"]), fmt(r["reliability_index_at_peak"])] for r in mc],
        [2.15, 0.85, 0.95, 1.0, 1.0, 1.0],
    )
    p(
        doc,
        "Table 4 quantifies the state effect. The cumulative failure probability is 0.157 for the bare slope, 0.669 for the cracked slope, 0.043 for the cracked vegetated slope, 0.883 for the road-surcharge/clogged-drain state and 0.017 for the road-surcharge/maintained-drainage state. These values should not be interpreted as site-specific probabilities; they are reproducible indicators showing how the same uncertain slope can change risk class when crack, vegetation and infrastructure states are modified.",
    )
    p(
        doc,
        "Figure 4 gives the rainfall threshold envelope and includes the Caine (1980) empirical lower-bound threshold as an external reference. The empirical line is not used as calibration because it summarizes heterogeneous landslide observations, but it anchors the scenario-conditioned curves against a widely cited rainfall-duration criterion. Instead of a single empirical intensity-duration line, the proposed envelope changes with slope state. This is the central operational implication: a rainfall event acceptable for a vegetated slope with maintained drainage may exceed the stability threshold for a cracked slope or for a road cut with clogged drainage.",
    )
    caption(doc, "Fig. 4 State-conditioned rainfall threshold envelope for deterministic FS=1.")
    doc.add_picture(str(FIG / "Fig4_threshold_envelope.png"), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p(
        doc,
        "Sensitivity analysis in Figure 5 shows that the cracked-slope minimum factor of safety is controlled mainly by crack depth, friction angle, cohesion, initial suction and hydraulic conductivity. This ranking is consistent with the physical mechanism: crack depth and hydraulic conductivity control wetting and crack water pressure, while friction and cohesion control the resistance margin (Cho 2007, 2014).",
    )
    caption(doc, "Fig. 5 Rank-based sensitivity controls for minimum FS in the cracked-slope state.")
    doc.add_picture(str(FIG / "Fig5_sensitivity_controls.png"), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p(
        doc,
        "Finally, Figure 6 summarizes the risk contrast by state. The road-surcharge/clogged-drain state dominates the risk space, while the maintained-drainage state with vegetation produces a probability lower than the bare slope. This result supports a practical engineering-geology message: maintenance state can be treated as a formal reliability variable, not merely as a qualitative site note.",
    )
    caption(doc, "Fig. 6 Risk contrast by slope state, combining cumulative failure probability and median FS deficit.")
    doc.add_picture(str(FIG / "Fig6_state_risk_contrast.png"), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    metrics = read_csv("proposed_metrics.csv")
    caption(doc, "Table 5 Proposed decision metrics and interpretation.")
    plain_table(
        doc,
        ["Metric", "Definition", "Interpretation"],
        [[r["metric"], r["definition"], r["interpretation"]] for r in metrics],
        [1.55, 2.7, 2.15],
    )
    p(
        doc,
        "The CRCF computed from Table 4 is approximately 1.22, meaning that the selected vegetation state more than neutralizes the crack-induced risk increment relative to the bare slope in this synthetic example. The infrastructure drainage resilience is approximately 0.98, meaning that the maintained-drainage state reduces most of the risk associated with the road/clogged-drain condition. These values are scenario-specific, but the definitions are general and can be recalculated for site-specific data.",
    )
    p(
        doc,
        "Figure 7 checks Monte Carlo convergence for the most decision-sensitive states. The mean probability estimates stabilize as the sample size increases from 1,000 to 30,000 samples, and the replicate standard deviation becomes small relative to the state contrast. This check does not turn the synthetic case into field validation, but it reduces the risk that the reported state ranking is a sampling artifact.",
    )
    caption(doc, "Fig. 7 Monte Carlo convergence check for selected high- and low-risk slope states.")
    doc.add_picture(str(FIG / "Fig7_monte_carlo_convergence.png"), width=Inches(6.2))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()
    caption(doc, "Table 6 Comparison between the proposed framework and common assessment alternatives.")
    plain_table(
        doc,
        ["Approach", "Typical output", "Limitation addressed by the proposed framework"],
        [
            ["Static factor of safety", "Single FS", "Cannot represent storm timing or post-rainfall memory"],
            ["Transient seepage plus FS", "FS(t)", "Usually omits explicit crack-root-infrastructure state comparison"],
            ["Empirical intensity-duration threshold", "Rainfall trigger line", "Does not encode slope biological or drainage maintenance state"],
            ["Conventional Monte Carlo slope reliability", "Pf for one model state", "Does not separate risk changes caused by cracks, roots and infrastructure"],
            ["Proposed state-conditioned framework", "FSmin(t), Pf(t), HMMI, CRCF, IDR and threshold envelope", "Links hydrological forcing, biological reinforcement and civil-infrastructure maintenance in one reproducible decision space"],
        ],
        [1.7, 1.6, 3.1],
    )

    heading(doc, "5 Discussion", 1)
    p(
        doc,
        "The proposed framework is intentionally modest in its hydraulic mechanics and ambitious in its decision structure. A full Richards solution with spatially variable boundary conditions would be preferable for final design, but it can obscure the state comparison when only screening-level data are available. The reduced wetness variable makes the assumptions visible and reproducible, which is useful during early engineering-geology appraisal and during maintenance prioritization. The external comparison with the Caine (1980) lower-bound threshold and the comparison in Table 6 are not substitutes for site calibration, but they locate the proposed outputs relative to established empirical and deterministic practice.",
    )
    p(
        doc,
        "The infrastructure extension is the main state-of-practice contribution. Road surcharges, surface drains and culverts are usually described qualitatively during field reconnaissance, but they are not always inserted into probabilistic rainfall-stability calculations. By treating surcharge, runoff concentration and drainage dissipation as explicit state variables, the method converts maintenance into a reliability input. This enables comparison of vegetation, drainage cleaning and structural loading under the same rainfall envelope.",
    )
    p(
        doc,
        "The framework also clarifies the role of vegetation. Root reinforcement is not presented as a universal stabilizer. It is effective only when root depth, root cohesion and hydrological interception interact with the failure depth and storm timing. This is consistent with bioengineering literature, where mechanical and hydrological root effects depend on both soil and plant state (Wu et al. 1979; Gray and Sotir 1992; Schwarz et al. 2010; Sidle and Bogaard 2016).",
    )
    p(
        doc,
        "Several limitations must be kept explicit. The demonstration uses a synthetic slope and reduced hydrology, not field monitoring. The crack hydraulic load is simplified, the road surcharge is uniform, and root cohesion is represented as an apparent cohesion rather than a mechanistic root-distribution model. The results therefore support method evaluation and reproducibility, not direct design for a particular road cut. Future work should connect the same metrics to monitored slopes, spatially distributed infiltration, unsaturated finite elements and observed maintenance histories.",
    )

    heading(doc, "6 Conclusions", 1)
    p(
        doc,
        "This paper presented a reproducible state-conditioned framework for rainfall-induced stability assessment of cracked, vegetated and infrastructure-modified slopes. The framework replaces a single static factor of safety with FSmin(t), Pf(t), HMMI, CRCF, infrastructure drainage resilience and a state-conditioned rainfall threshold envelope. These outputs are intended to make transient hydrological memory and civil-infrastructure condition visible during engineering-geology decisions.",
    )
    p(
        doc,
        "The numerical demonstration showed that crack state can shift the synthetic slope from moderate probability to high probability of failure, while vegetation can strongly reduce risk when its root and hydrological effects intersect the critical zone. Adding road surcharge and clogged drainage produced the highest risk state, whereas maintained drainage and vegetation reduced the risk below the bare-slope condition. The central conclusion is that drainage maintenance and biological reinforcement can be assessed as quantitative reliability states rather than as qualitative modifiers.",
    )
    p(
        doc,
        "The proposed metrics are not substitutes for site investigation or detailed design. Their value is to organize transparent comparisons, expose hidden assumptions and connect rainfall thresholds to the actual physical state of the slope and nearby infrastructure. Because the code, random seed, tables and figures are supplied, reviewers can reproduce or replace the synthetic demonstration with site-specific parameters.",
    )

    heading(doc, "Data and code availability", 1)
    p(
        doc,
        f"The data used in this study are synthetic and were generated by the reproducible code supplied with the manuscript and mirrored in a public repository: {GITHUB_URL}. The computational package includes the Python script, fixed random seed, CSV outputs, figure files and requirements file. No field or personal data were used.",
    )
    heading(doc, "Statements and Declarations", 1)
    p(doc, "Competing interests: The author declares no competing interests.")
    p(doc, "Funding: This research received no external funding.")
    p(
        doc,
        "Author contributions: G.J.M.C. conceived the study, developed the methodology, prepared the numerical demonstration, generated the figures and tables, wrote the manuscript and reviewed the final version.",
    )
    p(
        doc,
        "Declaration of generative AI and AI-assisted technologies: During the preparation of this submission version, a generative AI-assisted coding and editing tool was used to assist with journal-format adaptation, English language restructuring, reproducible calculation preparation and checklist review. The author reviewed and edited the content and takes full responsibility for the submitted manuscript.",
    )

    heading(doc, "References", 1)
    refs = references()
    for ref in refs:
        reference_paragraph(doc, ref)

    doc.save(OUT_DOCX)
    return OUT_DOCX


def references() -> list[str]:
    return [
        "Abramson LW, Lee TS, Sharma S, Boyce GM (2002) Slope stability and stabilization methods, 2nd edn. Wiley, New York.",
        "Ang AHS, Tang WH (2007) Probability concepts in engineering: emphasis on applications to civil and environmental engineering, 2nd edn. Wiley, New York.",
        "Baecher GB, Christian JT (2003) Reliability and statistics in geotechnical engineering. Wiley, Chichester.",
        "Baum RL, Savage WZ, Godt JW (2008) TRIGRS - A Fortran program for transient rainfall infiltration and grid-based regional slope-stability analysis, version 2.0. U.S. Geological Survey Open-File Report 2008-1159. https://doi.org/10.3133/ofr20081159",
        "Caine N (1980) The rainfall intensity-duration control of shallow landslides and debris flows. Geografiska Annaler: Series A, Physical Geography 62:23-27. https://doi.org/10.1080/04353676.1980.11879996",
        "Cho SE (2007) Effects of spatial variability of soil properties on slope stability. Engineering Geology 92:97-109. https://doi.org/10.1016/j.enggeo.2007.03.006",
        "Cho SE (2014) Probabilistic stability analysis of rainfall-induced landslides considering spatial variability of permeability. Engineering Geology 171:11-20. https://doi.org/10.1016/j.enggeo.2013.12.015",
        "Collins BD, Znidarcic D (2004) Stability analyses of rainfall induced landslides. Journal of Geotechnical and Geoenvironmental Engineering 130:362-372. https://doi.org/10.1061/(ASCE)1090-0241(2004)130:4(362)",
        "Ditlevsen O, Madsen HO (1996) Structural reliability methods. Wiley, Chichester.",
        "Duncan JM (2000) Factors of safety and reliability in geotechnical engineering. Journal of Geotechnical and Geoenvironmental Engineering 126:307-316. https://doi.org/10.1061/(ASCE)1090-0241(2000)126:4(307)",
        "Duncan JM, Wright SG, Brandon TL (2014) Soil strength and slope stability, 2nd edn. Wiley, Hoboken.",
        "Fenton GA, Griffiths DV (2008) Risk assessment in geotechnical engineering. Wiley, Hoboken. https://doi.org/10.1002/9780470284704",
        "Fredlund DG, Morgenstern NR, Widger RA (1978) The shear strength of unsaturated soils. Canadian Geotechnical Journal 15:313-321. https://doi.org/10.1139/t78-029",
        "Fredlund DG, Rahardjo H (1993) Soil mechanics for unsaturated soils. Wiley, New York. https://doi.org/10.1002/9780470172759",
        "Fredlund DG, Rahardjo H, Fredlund MD (2012) Unsaturated soil mechanics in engineering practice. Wiley, Hoboken. https://doi.org/10.1002/9781118280492",
        "Gray DH, Sotir RB (1992) Biotechnical stabilization of highway cut slope. Journal of Geotechnical Engineering 118:1395-1409. https://doi.org/10.1061/(ASCE)0733-9410(1992)118:9(1395)",
        "Griffiths DV, Lu N (2005) Unsaturated slope stability analysis with steady infiltration or evaporation using elasto-plastic finite elements. International Journal for Numerical and Analytical Methods in Geomechanics 29:249-267.",
        "Guzzetti F, Peruccacci S, Rossi M, Stark CP (2008) The rainfall intensity-duration control of shallow landslides and debris flows: an update. Landslides 5:3-17. https://doi.org/10.1007/s10346-007-0112-1",
        "Iverson RM (2000) Landslide triggering by rain infiltration. Water Resources Research 36:1897-1910. https://doi.org/10.1029/2000WR900090",
        "Melchers RE, Beck AT (2018) Structural reliability analysis and prediction, 3rd edn. Wiley, Hoboken. https://doi.org/10.1002/9781119266105",
        "Montgomery DR, Dietrich WE (1994) A physically based model for the topographic control on shallow landsliding. Water Resources Research 30:1153-1171. https://doi.org/10.1029/93WR02979",
        "Morgan RPC, Rickson RJ (2003) Slope stabilization and erosion control: a bioengineering approach. Taylor & Francis, London. https://doi.org/10.4324/9780203362136",
        "Mualem Y (1976) A new model for predicting the hydraulic conductivity of unsaturated porous media. Water Resources Research 12:513-522. https://doi.org/10.1029/WR012i003p00513",
        "Ng CWW, Shi Q (1998) A numerical investigation of the stability of unsaturated soil slopes subjected to transient seepage. Computers and Geotechnics 22:1-28. https://doi.org/10.1016/S0266-352X(97)00036-0",
        "Phoon KK, Kulhawy FH (1999) Characterization of geotechnical variability. Canadian Geotechnical Journal 36:612-624. https://doi.org/10.1139/t99-038",
        "Rahardjo H, Leong EC, Rezaur RB (2008) Effect of antecedent rainfall on pore-water pressure distribution characteristics in residual soil slopes under tropical rainfall. Hydrological Processes 22:506-523. https://doi.org/10.1002/hyp.6880",
        "Richards LA (1931) Capillary conduction of liquids through porous mediums. Physics 1:318-333. https://doi.org/10.1063/1.1745010",
        "Schwarz M, Preti F, Giadrossich F, Lehmann P, Or D (2010) Quantifying the role of vegetation in slope stability: a case study in Tuscany, Italy. Ecological Engineering 36:285-291. https://doi.org/10.1016/j.ecoleng.2009.06.014",
        "Sidle RC, Bogaard TA (2016) Dynamic earth system and ecological controls of rainfall-initiated landslides. Earth-Science Reviews 159:275-291. https://doi.org/10.1016/j.earscirev.2016.05.013",
        "van Genuchten MT (1980) A closed-form equation for predicting the hydraulic conductivity of unsaturated soils. Soil Science Society of America Journal 44:892-898. https://doi.org/10.2136/sssaj1980.03615995004400050002x",
        "Waldron LJ, Dakessian S (1981) Soil reinforcement by roots. Soil Science 132:427-435. https://doi.org/10.1097/00010694-198112000-00007",
        "Wu TH, McKinnell WP III, Swanston DN (1979) Strength of tree roots and landslides on Prince of Wales Island, Alaska. Canadian Geotechnical Journal 16:19-33. https://doi.org/10.1139/t79-003",
    ]


def build_cover_letter():
    doc = Document()
    configure(doc)
    p(doc, "Dear Editor,", None)
    p(
        doc,
        f"Please consider the manuscript entitled \"Time-dependent stability of rainfall-infiltrated cracked, vegetated and infrastructure-modified slopes\" for publication in {JOURNAL}.",
    )
    p(
        doc,
        "The manuscript proposes a reproducible state-conditioned framework for rainfall-induced slope stability that explicitly combines crack hydraulics, vegetation effects and civil-infrastructure conditions such as road surcharge and drainage maintenance. The paper contributes hydro-mechanical memory, crack-root compensation and infrastructure drainage-resilience metrics, supported by a transparent Python demonstration with 30,000 Monte Carlo samples.",
    )
    p(
        doc,
        "The work has not been published previously and is not under consideration elsewhere. The data are synthetic and fully reproducible from the supplied computational package. The author declares no competing interests and no external funding.",
    )
    p(doc, "Sincerely,")
    p(doc, AUTHOR)
    p(doc, "Universidad de Panamá")
    p(doc, f"Email: {EMAIL}")
    p(doc, "Telephone: +507 6719-0245")
    doc.save(COVER)


def build_declarations():
    doc = Document()
    configure(doc)
    heading(doc, "Submission Declarations", 1)
    p(doc, "Manuscript title: Time-dependent stability of rainfall-infiltrated cracked, vegetated and infrastructure-modified slopes")
    p(doc, f"Target journal: {JOURNAL}")
    p(doc, f"Author: {AUTHOR}")
    heading(doc, "Competing Interests", 2)
    p(doc, "The author declares no competing interests.")
    heading(doc, "Funding", 2)
    p(doc, "This research received no external funding.")
    heading(doc, "Data Availability", 2)
    p(doc, f"The data used in this study are synthetic and were generated by the reproducible code supplied with the manuscript and mirrored in a public repository: {GITHUB_URL}. The computational package includes the Python script, fixed random seed, CSV outputs, figures and requirements file. No field or personal data were used.")
    heading(doc, "Declaration of Generative AI and AI-Assisted Technologies", 2)
    p(doc, "During the preparation of this submission version, a generative AI-assisted coding and editing tool was used to assist with journal-format adaptation, English language restructuring, reproducible calculation preparation and checklist review. The author reviewed and edited the content and takes full responsibility for the submitted manuscript.")
    heading(doc, "Author Information", 2)
    p(doc, AUTHOR)
    p(doc, AFFIL)
    p(doc, f"Corresponding author e-mail: {EMAIL}")
    p(doc, f"ORCID: {ORCID}")
    doc.save(DECL)


if __name__ == "__main__":
    build_manuscript()
    build_cover_letter()
    build_declarations()
